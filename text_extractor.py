import PyPDF2
import docx
import os
import sys
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF file
    
    Args:
        file_path (str): Path to PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            logger.info(f"PDF has {len(pdf_reader.pages)} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            logger.info(f"Total extracted text length: {len(text)} characters")
            return text
            
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file
    
    Args:
        file_path (str): Path to DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            text += para_text + "\n"
            if para_text.strip():  # Only log non-empty paragraphs
                logger.info(f"Paragraph {para_num + 1}: {len(para_text)} characters")
        
        # Extract text from tables if any
        for table_num, table in enumerate(doc.tables):
            logger.info(f"Processing table {table_num + 1}")
            for row_num, row in enumerate(table.rows):
                for cell_num, cell in enumerate(row.cells):
                    cell_text = cell.text
                    text += cell_text + " "
                text += "\n"
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(file_path):
    """
    Extract text from TXT file
    
    Args:
        file_path (str): Path to TXT file
        
    Returns:
        str: Extracted text content
    """
    try:
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    logger.info(f"Text length: {len(text)} characters")
                    return text
            except UnicodeDecodeError:
                logger.warning(f"Failed to read with {encoding} encoding, trying next...")
                continue
        
        logger.error("Failed to read TXT file with any supported encoding")
        return ""
        
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        return ""

def extract_text_from_file(file_path):
    """
    Extract text from file based on extension
    
    Args:
        file_path (str): Path to file
        
    Returns:
        tuple: (extracted_text, file_info)
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return "", {}
    
    file_path = Path(file_path)
    file_extension = file_path.suffix.lower()
    file_size = os.path.getsize(file_path)
    
    file_info = {
        'filename': file_path.name,
        'extension': file_extension,
        'size_bytes': file_size,
        'size_mb': round(file_size / (1024 * 1024), 2)
    }
    
    logger.info(f"Processing file: {file_info}")
    
    extracted_text = ""
    
    if file_extension == '.pdf':
        extracted_text = extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        extracted_text = extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        extracted_text = extract_text_from_txt(file_path)
    else:
        logger.error(f"Unsupported file format: {file_extension}")
        return "", file_info
    
    # Generate content preview (first 200 characters)
    preview = extracted_text[:200].replace('\n', ' ').strip()
    if len(extracted_text) > 200:
        preview += "..."
    
    file_info['content_preview'] = preview
    file_info['total_characters'] = len(extracted_text)
    file_info['total_words'] = len(extracted_text.split()) if extracted_text else 0
    
    return extracted_text, file_info

def test_extraction_from_directory(directory_path="uploads"):
    """
    Test text extraction from all files in uploads directory
    
    Args:
        directory_path (str): Path to directory containing files to test
    """
    if not os.path.exists(directory_path):
        logger.error(f"Directory not found: {directory_path}")
        return
    
    supported_extensions = ['.pdf', '.docx', '.txt']
    files_found = []
    
    for file_path in Path(directory_path).rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            files_found.append(file_path)
    
    if not files_found:
        logger.info(f"No supported files found in {directory_path}")
        logger.info(f"Supported formats: {supported_extensions}")
        return
    
    logger.info(f"Found {len(files_found)} files to test")
    
    for file_path in files_found:
        logger.info(f"\n{'='*50}")
        logger.info(f"Testing file: {file_path}")
        logger.info(f"{'='*50}")
        
        extracted_text, file_info = extract_text_from_file(file_path)
        
        if extracted_text:
            print(f"\n✅ SUCCESS: {file_info['filename']}")
            print(f"   📄 Size: {file_info['size_mb']} MB")
            print(f"   📝 Characters: {file_info['total_characters']}")
            print(f"   📊 Words: {file_info['total_words']}")
            print(f"   👀 Preview: {file_info['content_preview']}")
        else:
            print(f"\n❌ FAILED: {file_info['filename']}")
            print(f"   📄 Size: {file_info['size_mb']} MB")

def test_single_file(file_path):
    """
    Test text extraction from a single file
    
    Args:
        file_path (str): Path to file to test
    """
    logger.info(f"Testing single file: {file_path}")
    
    extracted_text, file_info = extract_text_from_file(file_path)
    
    print(f"\n{'='*60}")
    print(f"FILE EXTRACTION TEST RESULTS")
    print(f"{'='*60}")
    print(f"File: {file_info.get('filename', 'Unknown')}")
    print(f"Extension: {file_info.get('extension', 'Unknown')}")
    print(f"Size: {file_info.get('size_mb', 0)} MB")
    
    if extracted_text:
        print(f"Status: ✅ SUCCESS")
        print(f"Characters extracted: {file_info.get('total_characters', 0)}")
        print(f"Words extracted: {file_info.get('total_words', 0)}")
        print(f"\nPreview:")
        print(f"{'-'*40}")
        print(file_info.get('content_preview', 'No preview available'))
        print(f"{'-'*40}")
        
        # Show first few lines of actual content
        lines = extracted_text.split('\n')[:5]
        print(f"\nFirst 5 lines:")
        for i, line in enumerate(lines, 1):
            if line.strip():
                print(f"{i}: {line.strip()[:100]}...")
    else:
        print(f"Status: ❌ FAILED")
        print("No text could be extracted from this file.")

if __name__ == "__main__":
    print("🔍 Document Text Extractor Testing Tool")
    print("="*50)
    
    if len(sys.argv) > 1:
        # Test specific file
        file_path = sys.argv[1]
        test_single_file(file_path)
    else:
        # Test all files in uploads directory
        print("Testing all files in 'uploads' directory...")
        test_extraction_from_directory()
        
        print(f"\n{'='*50}")
        print("💡 Usage tips:")
        print("   - To test a specific file: python text_extractor.py <file_path>")
        print("   - To test all files: python text_extractor.py")
        print("   - Supported formats: PDF, DOCX, TXT")
        print("   - Check logs above for detailed extraction info")